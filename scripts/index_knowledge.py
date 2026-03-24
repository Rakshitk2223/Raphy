import argparse
import sys
from pathlib import Path

import chromadb
from sentence_transformers import SentenceTransformer

try:
    import fitz  # pymupdf
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import markdownify
except ImportError:
    markdownify = None


CHUNK_SIZE = 256
CHUNK_OVERLAP = 30


def convert_pdf_to_markdown(file_path: Path) -> str:
    if fitz is None:
        raise ImportError("pymupdf not installed")

    doc = fitz.open(str(file_path))
    text_parts = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            text_parts.append(text)

    return "\n\n".join(text_parts)


def convert_docx_to_markdown(file_path: Path) -> str:
    if docx is None:
        raise ImportError("python-docx not installed")
    if markdownify is None:
        raise ImportError("markdownify not installed")

    doc = docx.Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        table_md = "| " + " | ".join(["Header"] * len(table.columns)) + " |\n"
        table_md += "| " + " | ".join(["---"] * len(table.columns)) + " |\n"

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md += "| " + " | ".join(cells) + " |\n"

        parts.append(table_md)

    return "\n\n".join(parts)


def convert_file_to_markdown(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return convert_pdf_to_markdown(file_path)
    elif suffix in [".docx", ".doc"]:
        return convert_docx_to_markdown(file_path)
    elif suffix in [".txt", ".md", ".markdown"]:
        return file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    words = text.split()
    chunks = []

    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)

    return chunks


def get_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")


def main():
    parser = argparse.ArgumentParser(description="Index knowledge files to vector DB")
    parser.add_argument(
        "--knowledge-dir",
        type=Path,
        default=Path("data/memory/knowledge"),
        help="Directory containing files to index",
    )
    parser.add_argument(
        "--chroma-dir",
        type=Path,
        default=Path("data/memory/chroma"),
        help="Directory to store ChromaDB",
    )
    parser.add_argument("--chunk-size", type=int, default=CHUNK_SIZE, help="Chunk size in tokens")
    parser.add_argument(
        "--chunk-overlap", type=int, default=CHUNK_OVERLAP, help="Overlap between chunks"
    )
    parser.add_argument(
        "--clear", action="store_true", help="Clear existing collection before indexing"
    )

    args = parser.parse_args()

    knowledge_dir = args.knowledge_dir
    chroma_dir = args.chroma_dir

    if not knowledge_dir.exists():
        print(f"Knowledge directory does not exist: {knowledge_dir}")
        print(f"Create it and add files to index: mkdir -p {knowledge_dir}")
        sys.exit(1)

    supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}
    files = [f for f in knowledge_dir.iterdir() if f.suffix.lower() in supported_extensions]

    if not files:
        print(f"No supported files found in {knowledge_dir}")
        print(f"Supported: {', '.join(supported_extensions)}")
        sys.exit(1)

    print(f"Found {len(files)} files to index")

    chroma_dir.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(chroma_dir))

    collection = client.get_or_create_collection("knowledge")

    if args.clear:
        print("Clearing existing collection...")
        client.delete_collection("knowledge")
        collection = client.get_or_create_collection("knowledge")

    model = get_embedding_model()
    print(f"Using embedding model: {model}")

    total_chunks = 0

    for file_path in files:
        print(f"\nProcessing: {file_path.name}")

        try:
            markdown_text = convert_file_to_markdown(file_path)
            print(f"  Converted to markdown ({len(markdown_text)} chars)")
        except Exception as e:
            print(f"  ERROR converting file: {e}")
            continue

        chunks = chunk_text(markdown_text, args.chunk_size, args.chunk_overlap)
        print(f"  Created {len(chunks)} chunks")

        ids = []
        documents = []
        metadatas = []

        for i, chunk in enumerate(chunks):
            chunk_id = f"{file_path.stem}_{i}"
            ids.append(chunk_id)
            documents.append(chunk)
            metadatas.append({"source": file_path.name, "chunk_index": i})

        if ids:
            embeddings = model.encode(documents).tolist()
            collection.add(ids=ids, documents=documents, embeddings=embeddings, metadatas=metadatas)
            total_chunks += len(chunks)
            print(f"  Indexed {len(chunks)} chunks")

    print(f"\n✓ Indexed {total_chunks} chunks from {len(files)} files")
    print(f"  Stored in: {chroma_dir}")


if __name__ == "__main__":
    main()
